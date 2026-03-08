import streamlit as st
from docxtpl import DocxTemplate
from docx import Document
import io, re, json, requests, base64, mammoth
from datetime import datetime

# --- 1. CẤU HÌNH GITHUB (mrspytro) ---
GITHUB_USER = "mrspytro"
GITHUB_REPO = "Form_Auto"
GITHUB_FOLDER = "templates"

# Token fen cung cấp: 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
# KHUYÊN DÙNG: Dán vào mục Secrets của Streamlit Cloud với tên GITHUB_TOKEN
GITHUB_TOKEN = st.secrets.get("GITHUB_TOKEN", "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")

# Header chuẩn cho Fine-grained Token (Cần Bearer và API Version)
HEADERS = {
    "Authorization": f"Bearer {GITHUB_TOKEN.strip()}",
    "Accept": "application/vnd.github+json",
    "X-GitHub-Api-Version": "2022-11-28"
}

# --- 2. HÀM XỬ LÝ GITHUB API ---
def get_online_templates():
    url = f"https://api.github.com/repos/{GITHUB_USER}/{GITHUB_REPO}/contents/{GITHUB_FOLDER}"
    try:
        res = requests.get(url, headers=HEADERS)
        if res.status_code == 200:
            return {f['name']: f['download_url'] for f in res.json() if f['name'].endswith('.docx')}
        elif res.status_code == 401:
            st.sidebar.error("Lỗi 401: Token không hợp lệ hoặc thiếu quyền (Bad credentials).")
        return {}
    except: return {}

def upload_to_github(file_name, file_bytes):
    url = f"https://api.github.com/repos/{GITHUB_USER}/{GITHUB_REPO}/contents/{GITHUB_FOLDER}/{file_name}"
    encoded_content = base64.b64encode(file_bytes).decode("utf-8")
    data = {"message": f"Update mẫu: {file_name}", "content": encoded_content, "branch": "main"}
    
    # Check SHA để ghi đè nếu file đã tồn tại
    check_res = requests.get(url, headers=HEADERS)
    if check_res.status_code == 200:
        data["sha"] = check_res.json()["sha"]

    res = requests.put(url, headers=HEADERS, json=data)
    if res.status_code in [200, 201]: return True, "Thành công"
    return False, f"Lỗi {res.status_code}: {res.json().get('message')}"

# --- 3. LOGIC XỬ LÝ BIẾN WORD (Sắp xếp theo thứ tự văn bản) ---
def get_variables_v13(file_stream):
    try:
        doc = Document(file_stream)
        ordered_vars, seen = [], set()
        pattern = r"\{\{\s*(\w+)\s*\}\}"
        
        def process_match(match_text):
            m = re.match(r"([a-zA-Z]+)(\d+)_((?:(?!__).)+)(?:__(.*))?", match_text)
            if m:
                prefix, num, raw_name, note = m.groups()
                p_lower = prefix.lower()
                v_type = 'title' if p_lower == 't' else ('checkbox' if p_lower == 'cb' else 'field')
                display_name = raw_name.replace("_", " ").strip()
                return {
                    'original': match_text, 'type': v_type,
                    'label': display_name.upper() if v_type == 'title' else f"Mục {num}: {display_name.title()}",
                    'note': (note.replace("_", " ") if note else "")
                }
            return None

        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = [p for p in doc.paragraphs if p._element == element]
                if para:
                    for var in re.findall(pattern, para[0].text):
                        if var not in seen:
                            res = process_match(var)
                            if res: ordered_vars.append(res); seen.add(var)
            elif element.tag.endswith('tbl'):
                table = [t for t in doc.tables if t._element == element]
                if table:
                    for row in table[0].rows:
                        for cell in row.cells:
                            for var in re.findall(pattern, cell.text):
                                if var not in seen:
                                    res = process_match(var);
                                    if res: ordered_vars.append(res); seen.add(var)
        return ordered_vars
    except: return []

# --- 4. GIAO DIỆN XÁC THỰC ---
def login():
    if "auth" not in st.session_state: st.session_state.auth = False
    if st.session_state.auth: return True
    st.title("🔐 Hệ thống mrspytro 1992")
    pwd = st.text_input("Mật khẩu truy cập", type="password")
    if st.button("Đăng nhập"):
        if pwd == "phi1992": st.session_state.auth = True; st.rerun()
        else: st.error("Mật khẩu không đúng!")
    return False

# --- 5. GIAO DIỆN CHÍNH ---
if login():
    st.set_page_config(page_title="Form Auto Pro", layout="wide")
    
    st.markdown("""
        <style>
        .cat-box { background: #f8f9fa; padding: 12px; border-left: 6px solid #d32f2f; margin: 20px 0; border-radius: 4px; }
        .paper-view { background: #525659; padding: 30px 0; display: flex; justify-content: center; }
        .a4-page { background: white; width: 210mm; min-height: 297mm; padding: 20mm; box-shadow: 0 0 15px rgba(0,0,0,0.5); color: black; font-family: 'Times New Roman', serif; overflow: hidden; }
        .a4-page table { border-collapse: collapse; width: 100% !important; }
        .a4-page td { border: 1px solid black; padding: 5px; }
        </style>
        """, unsafe_allow_html=True)

    if 'form_data' not in st.session_state: st.session_state.form_data = {}

    with st.sidebar:
        st.header("📂 Kho mẫu hồ sơ")
        templates = get_online_templates()
        if templates:
            selected = st.selectbox("Chọn mẫu hiện có:", ["--- Chọn ---"] + list(templates.keys()))
            active_url = templates.get(selected)
        else:
            st.warning("Đang quét thư mục /templates...")
            active_url = None

        st.divider()
        st.header("📤 Lưu mẫu mới")
        new_file = st.file_uploader("Upload .docx lên GitHub", type=["docx"])
        if new_file and st.button("Lưu lên GitHub"):
            ok, msg = upload_to_github(new_file.name, new_file.getvalue())
            if ok: st.success("Đã lưu thành công!"); st.rerun()
            else: st.error(msg)
        
        if st.button("🗑️ Xóa sạch Form"): st.session_state.form_data = {}; st.rerun()

    # --- QUY TRÌNH XỬ LÝ BIẾN ---
    file_bytes = None
    if active_url:
        try:
            res = requests.get(active_url)
            file_bytes = io.BytesIO(res.content)
        except: st.error("Không thể tải file mẫu.")

    if file_bytes:
        vars_list = get_variables_v13(io.BytesIO(file_bytes.getvalue()))
        tab_in, tab_pre = st.tabs(["📝 NHẬP LIỆU", "🔍 XEM TRƯỚC"])

        with tab_in:
            now = datetime.now()
            for item in vars_list:
                k = item['original']
                if item['type'] == 'title':
                    st.markdown(f'<div class="cat-box">### 📂 {item["label"]}</div>', unsafe_allow_html=True)
                    st.session_state.form_data[k] = ""
                else:
                    c1, c2 = st.columns([0.65, 0.35])
                    with c1:
                        # Auto-fill Ngày/Tháng/Năm
                        lbl = item['label'].lower()
                        d_val = ""
                        if "ngay" in lbl: d_val = str(now.day)
                        elif "thang" in lbl: d_val = str(now.month)
                        elif "nam" in lbl: d_val = str(now.year)
                        
                        v = st.session_state.form_data.get(k, False if item['type'] == 'checkbox' else d_val)
                        
                        if item['type'] == 'checkbox':
                            st.session_state.form_data[k] = st.checkbox(item['label'], value=v, key=f"c_{k}")
                        else:
                            st.session_state.form_data[k] = st.text_input(item['label'], value=v, key=f"t_{k}")
                    with c2:
                        if item['note']: st.caption(f"💡 *{item['note']}*")

        with tab_pre:
            # Render bản cuối
            ctx = {k: ("☑" if v is True else "☐" if v is False else v) for k, v in st.session_state.form_data.items()}
            try:
                doc = DocxTemplate(io.BytesIO(file_bytes.getvalue()))
                doc.render(ctx)
                out = io.BytesIO(); doc.save(out); out.seek(0)
                
                # Mammoth Preview
                html = mammoth.convert_to_html(out)
                st.markdown(f'<div class="paper-view"><div class="a4-page">{html.value}</div></div>', unsafe_allow_html=True)
                
                st.divider()
                st.download_button("🚀 TẢI FILE WORD HOÀN TẤT", data=out.getvalue(), file_name=f"Done_{selected}")
            except Exception as e: st.error(f"Lỗi Render: {e}")
    else:
        st.info("👈 Hãy chọn mẫu từ thư mục online hoặc tải mẫu mới lên để bắt đầu.")